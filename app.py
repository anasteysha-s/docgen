"""
DocGen - web service for automatic document generation from templates.
Backend: Python / Flask
Supported formats: PDF, DOCX
Document types: Service Contract, Completion Act, GDPR Request
"""

import os
import io
import uuid
from datetime import datetime
from flask import Flask, request, jsonify, send_file, render_template, render_template_string

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "generated_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ----------------------------------------------------------------------------
# FONT REGISTRATION (Cyrillic support)
# ----------------------------------------------------------------------------
# ReportLab's built-in Helvetica is WinAnsi-only, so it renders Cyrillic as
# black squares. We register a system TTF that covers Cyrillic and fall back
# to Helvetica only if nothing is found (Latin-only output).

PDF_FONT        = "Helvetica"
PDF_FONT_BOLD   = "Helvetica-Bold"
PDF_FONT_ITALIC = "Helvetica-Oblique"
DOCX_FONT       = "Arial"


def _register_pdf_font_pack(name, regular, bold, italic):
    if not (os.path.exists(regular) and os.path.exists(bold)):
        return False
    try:
        pdfmetrics.registerFont(TTFont(name,            regular))
        pdfmetrics.registerFont(TTFont(name + "-Bold",  bold))
        ital_path = italic if os.path.exists(italic) else regular
        pdfmetrics.registerFont(TTFont(name + "-Italic", ital_path))
        registerFontFamily(
            name,
            normal=name,
            bold=name + "-Bold",
            italic=name + "-Italic",
            boldItalic=name + "-Bold",
        )
        return True
    except Exception:
        return False


def _init_pdf_fonts():
    global PDF_FONT, PDF_FONT_BOLD, PDF_FONT_ITALIC
    here = os.path.dirname(__file__)
    candidates = [
        # Windows
        (r"C:\Windows\Fonts\arial.ttf",
         r"C:\Windows\Fonts\arialbd.ttf",
         r"C:\Windows\Fonts\ariali.ttf"),
        # Linux: DejaVu (most common)
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf"),
        # Linux: Liberation
        ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
         "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf"),
        # macOS
        ("/Library/Fonts/Arial.ttf",
         "/Library/Fonts/Arial Bold.ttf",
         "/Library/Fonts/Arial Italic.ttf"),
        # Bundled fallback
        (os.path.join(here, "fonts", "DejaVuSans.ttf"),
         os.path.join(here, "fonts", "DejaVuSans-Bold.ttf"),
         os.path.join(here, "fonts", "DejaVuSans-Oblique.ttf")),
    ]
    for reg, bold, ital in candidates:
        if _register_pdf_font_pack("DocSans", reg, bold, ital):
            PDF_FONT        = "DocSans"
            PDF_FONT_BOLD   = "DocSans-Bold"
            PDF_FONT_ITALIC = "DocSans-Italic"
            return


_init_pdf_fonts()

# ----------------------------------------------------------------------------
# DOCUMENT TEMPLATES
# ----------------------------------------------------------------------------

def get_contract_data(fields):
    return {
        "title": "ДОГОВІР ПРО НАДАННЯ ПОСЛУГ",
        "number": fields.get("doc_number", "001"),
        "city": fields.get("city", ""),
        "date": fields.get("date", datetime.today().strftime("%Y-%m-%d")),
        "sections": [
            {
                "heading": "1. ПРЕДМЕТ ДОГОВОРУ",
                "body": (
                    "1.1. Виконавець (" + fields.get("provider_name", "___") + ", надалі — «Виконавець») "
                    "зобов'язується надати Замовнику (" + fields.get("client_name", "___") + ", надалі — «Замовник») "
                    "такі послуги: " + fields.get("service_description", "___") + ".\n"
                    "1.2. Замовник зобов'язується прийняти та оплатити надані послуги відповідно до умов цього Договору."
                ),
            },
            {
                "heading": "2. ВАРТІСТЬ ПОСЛУГ ТА ПОРЯДОК РОЗРАХУНКІВ",
                "body": (
                    "2.1. Загальна вартість послуг становить " + fields.get("amount", "0") + " грн (з ПДВ).\n"
                    "2.2. Замовник здійснює оплату протягом " + fields.get("payment_days", "5") + " робочих днів "
                    "після підписання Акта виконаних робіт."
                ),
            },
            {
                "heading": "3. ТЕРМІНИ ВИКОНАННЯ",
                "body": (
                    "3.1. Виконавець зобов'язується надати послуги у строк до " + fields.get("deadline", "___") + ".\n"
                    "3.2. Дострокове виконання можливе за взаємною згодою Сторін."
                ),
            },
            {
                "heading": "4. ВІДПОВІДАЛЬНІСТЬ СТОРІН",
                "body": (
                    "4.1. За прострочення оплати Замовник сплачує пеню у розмірі 0,1% від суми "
                    "заборгованості за кожний день прострочення.\n"
                    "4.2. За невиконання або неналежне виконання послуг Виконавець несе відповідальність "
                    "відповідно до чинного законодавства."
                ),
            },
            {
                "heading": "5. РЕКВІЗИТИ СТОРІН",
                "body": (
                    "ВИКОНАВЕЦЬ: " + fields.get("provider_name", "___") + "\n"
                    "ЄДРПОУ / ІПН: " + fields.get("provider_code", "___") + "\n"
                    "Адреса: " + fields.get("provider_address", "___") + "\n\n"
                    "ЗАМОВНИК: " + fields.get("client_name", "___") + "\n"
                    "ЄДРПОУ / ІПН: " + fields.get("client_code", "___") + "\n"
                    "Адреса: " + fields.get("client_address", "___")
                ),
            },
        ],
    }


def _to_float(value):
    if value is None:
        return 0.0
    try:
        return float(str(value).replace(" ", "").replace(",", "."))
    except (ValueError, TypeError):
        return 0.0


def get_act_data(fields):
    services = fields.get("services") or []
    if not services:
        # Legacy single-service fallback (pre-multi-row API clients)
        services = [{
            "description": fields.get("service_description", ""),
            "unit":        fields.get("unit", ""),
            "quantity":    fields.get("quantity", "1"),
            "unit_price":  fields.get("unit_price", fields.get("amount", "0")),
        }]

    rows = []
    total = 0.0
    for i, s in enumerate(services, 1):
        qty   = _to_float(s.get("quantity", "1"))
        price = _to_float(s.get("unit_price", "0"))
        line_total = qty * price
        total += line_total
        rows.append([
            str(i),
            s.get("description") or "___",
            s.get("unit") or "service",
            s.get("quantity") or "1",
            "%.2f" % price,
            "%.2f" % line_total,
        ])

    total_str = "%.2f" % total

    return {
        "title": "АКТ ВИКОНАНИХ РОБІТ",
        "subtitle": "(Надання послуг)",
        "number": fields.get("doc_number", "001"),
        "city": fields.get("city", ""),
        "date": fields.get("date", datetime.today().strftime("%Y-%m-%d")),
        "sections": [
            {
                "heading": "ЗАГАЛЬНА ІНФОРМАЦІЯ",
                "body": (
                    "Виконавець: " + fields.get("provider_name", "___") + "\n"
                    "Замовник: " + fields.get("client_name", "___") + "\n"
                    "Підстава: Договір № " + fields.get("contract_number", "___") +
                    " від " + fields.get("contract_date", "___")
                ),
            },
            {
                "heading": "\nПЕРЕЛІК НАДАНИХ ПОСЛУГ",
                "table": {
                    "headers": ["№", "Найменування послуги", "Од.", "Кіл.", "Ціна, грн", "Сума, грн"],
                    "rows": rows,
                    "total": total_str,
                },
            },
            {
                "heading": "\nПІДСУМОК",
                "body": (
                    "Загальна вартість наданих послуг: " + total_str + " грн.\n"
                    "Послуги виконано в повному обсязі та у встановлені строки. "
                    "Претензій щодо якості та строків надання послуг Замовник не має."
                ),
            },
        ],
    }


def get_gdpr_data(fields):
    return {
        "title": "ЗАПИТ СУБ'ЄКТА ПЕРСОНАЛЬНИХ ДАНИХ",
        "subtitle": "(відповідно до Регламенту ЄС 2016/679 — GDPR)",
        "number": fields.get("doc_number", "001"),
        "city": fields.get("city", ""),
        "date": fields.get("date", datetime.today().strftime("%Y-%m-%d")),
        "sections": [
            {
                "heading": "1. ІДЕНТИФІКАЦІЯ СУБ'ЄКТА ДАНИХ",
                "body": (
                    "ПІБ: " + fields.get("client_name", "___") + "\n"
                    "Дата народження: " + fields.get("birth_date", "___") + "\n"
                    "Email: " + fields.get("email", "___") + "\n"
                    "Телефон: " + fields.get("phone", "___")
                ),
            },
            {
                "heading": "2. КОНТРОЛЕР ПЕРСОНАЛЬНИХ ДАНИХ",
                "body": (
                    "Організація: " + fields.get("provider_name", "___") + "\n"
                    "Адреса: " + fields.get("provider_address", "___") + "\n"
                    "Email DPO: " + fields.get("dpo_email", "___")
                ),
            },
            {
                "heading": "3. ВИД ЗАПИТУ",
                "body": (
                    "Суб'єкт даних просить реалізувати таке право: " +
                    fields.get("request_type", "право на доступ до персональних даних") + ".\n\n"
                    "Деталі: " + fields.get("request_details", "___")
                ),
            },
            {
                "heading": "4. ПРАВОВА ПІДСТАВА",
                "body": (
                    "Цей запит подається відповідно до:\n"
                    "- Статей 15–22 Регламенту ЄС 2016/679 (GDPR);\n"
                    "- Закону України «Про захист персональних даних».\n\n"
                    "Відповідно до статті 12 GDPR, прошу надати відповідь протягом 30 календарних днів."
                ),
            },
            {
                "heading": "5. ПІДТВЕРДЖЕННЯ",
                "body": (
                    "Я, " + fields.get("client_name", "___") + ", підтверджую достовірність наданих відомостей "
                    "та свідомо реалізую права суб'єкта персональних даних.\n\n"
                    "Дата: " + fields.get("date", datetime.today().strftime("%Y-%m-%d"))
                ),
            },
        ],
    }


DOCUMENT_BUILDERS = {
    "contract": get_contract_data,
    "act":      get_act_data,
    "gdpr":     get_gdpr_data,
}

# ----------------------------------------------------------------------------
# PDF GENERATION
# ----------------------------------------------------------------------------

def generate_pdf(doc_data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=2*cm, leftMargin=2.5*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("DocTitle", parent=styles["Normal"],
        fontSize=14, leading=18, alignment=TA_CENTER,
        spaceAfter=4, fontName=PDF_FONT_BOLD)
    subtitle_style = ParagraphStyle("DocSubtitle", parent=styles["Normal"],
        fontSize=9, leading=12, alignment=TA_CENTER, spaceAfter=2,
        fontName=PDF_FONT_ITALIC, textColor=colors.HexColor("#555555"))
    meta_style = ParagraphStyle("DocMeta", parent=styles["Normal"],
        fontSize=10, leading=14, alignment=TA_CENTER,
        spaceAfter=6, fontName=PDF_FONT)
    heading_style = ParagraphStyle("SectionHeading", parent=styles["Normal"],
        fontSize=11, leading=14, spaceBefore=14, spaceAfter=4,
        fontName=PDF_FONT_BOLD, textColor=colors.HexColor("#1a3a5c"))
    body_style = ParagraphStyle("BodyText", parent=styles["Normal"],
        fontSize=10, leading=15, alignment=TA_JUSTIFY,
        spaceAfter=4, fontName=PDF_FONT)
    sign_style = ParagraphStyle("SignLine", parent=styles["Normal"],
        fontSize=10, leading=20, fontName=PDF_FONT)

    story = []
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(doc_data["title"], title_style))
    if doc_data.get("subtitle"):
        story.append(Paragraph(doc_data["subtitle"], subtitle_style))
    story.append(Paragraph(
        "№ %s  |  %s  |  %s" % (doc_data["number"], doc_data["city"], doc_data["date"]),
        meta_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1a3a5c")))
    story.append(Spacer(1, 0.4*cm))

    for section in doc_data["sections"]:
        story.append(Paragraph(section["heading"], heading_style))
        if "body" in section:
            for line in section["body"].split("\n"):
                story.append(Paragraph(line or "&nbsp;", body_style))
        if "table" in section:
            t = section["table"]
            table_data = [t["headers"]] + t["rows"] + [["", "", "", "", "РАЗОМ:", t["total"]]]
            col_widths = [1*cm, 7*cm, 2*cm, 2*cm, 2.5*cm, 2.5*cm]
            tbl = Table(table_data, colWidths=col_widths)
            tbl.setStyle(TableStyle([
                ("FONTNAME",      (0,0),  (-1,-1), PDF_FONT),
                ("BACKGROUND",    (0,0),  (-1,0),  colors.HexColor("#1a3a5c")),
                ("TEXTCOLOR",     (0,0),  (-1,0),  colors.white),
                ("FONTNAME",      (0,0),  (-1,0),  PDF_FONT_BOLD),
                ("FONTSIZE",      (0,0),  (-1,-1), 9),
                ("ALIGN",         (0,0),  (-1,-1), "CENTER"),
                ("ALIGN",         (1,1),  (1,-2),  "LEFT"),
                ("GRID",          (0,0),  (-1,-2), 0.5, colors.HexColor("#cccccc")),
                ("BACKGROUND",    (0,-1), (-1,-1), colors.HexColor("#e8f0fe")),
                ("FONTNAME",      (0,-1), (-1,-1), PDF_FONT_BOLD),
                ("LINEABOVE",     (0,-1), (-1,-1), 1, colors.HexColor("#1a3a5c")),
                ("ROWBACKGROUNDS",(0,1),  (-1,-2), [colors.white, colors.HexColor("#f5f8ff")]),
                ("TOPPADDING",    (0,0),  (-1,-1), 5),
                ("BOTTOMPADDING", (0,0),  (-1,-1), 5),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 0.3*cm))

    story.append(Spacer(1, 0.8*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#aaaaaa")))
    story.append(Spacer(1, 0.4*cm))
    sign_data = [
        [Paragraph("ВИКОНАВЕЦЬ / ЗАЯВНИК", sign_style), Paragraph("ЗАМОВНИК / ОТРИМУВАЧ", sign_style)],
        [Paragraph("________________  /_____________/", sign_style), Paragraph("________________  /_____________/", sign_style)],
    ]
    sign_table = Table(sign_data, colWidths=[9*cm, 9*cm])
    sign_table.setStyle(TableStyle([
        ("ALIGN",      (0,0), (-1,-1), "LEFT"),
        ("FONTSIZE",   (0,0), (-1,-1), 10),
        ("TOPPADDING", (0,0), (-1,-1), 6),
    ]))
    story.append(sign_table)
    doc.build(story)
    buffer.seek(0)
    return buffer.read()


# ----------------------------------------------------------------------------
# DOCX GENERATION
# ----------------------------------------------------------------------------

def _set_run_font(run):
    """Apply Cyrillic-safe font to a run (ascii + hAnsi + cs)."""
    run.font.name = DOCX_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), DOCX_FONT)
    rFonts.set(qn("w:hAnsi"), DOCX_FONT)
    rFonts.set(qn("w:cs"),    DOCX_FONT)
    return run


def generate_docx(doc_data):
    document = Document()
    for section in document.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    # Set default font on the Normal style so unstyled runs inherit it.
    document.styles["Normal"].font.name = DOCX_FONT

    DARK_BLUE = RGBColor(0x1a, 0x3a, 0x5c)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_data["title"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    _set_run_font(run)

    if doc_data.get("subtitle"):
        p2 = document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(doc_data["subtitle"])
        r2.italic = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _set_run_font(r2)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("№ %s  |  %s  |  %s" % (doc_data["number"], doc_data["city"], doc_data["date"]))
    r3.font.size = Pt(10)
    _set_run_font(r3)

    document.add_paragraph()

    for section in doc_data["sections"]:
        ph = document.add_paragraph()
        rh = ph.add_run(section["heading"])
        rh.bold = True
        rh.font.size = Pt(11)
        rh.font.color.rgb = DARK_BLUE
        _set_run_font(rh)
        ph.space_before = Pt(12)

        if "body" in section:
            for line in section["body"].split("\n"):
                pb = document.add_paragraph()
                rb = pb.add_run(line)
                rb.font.size = Pt(10)
                _set_run_font(rb)
                pb.paragraph_format.space_after = Pt(2)

        if "table" in section:
            t = section["table"]
            rows = [t["headers"]] + t["rows"] + [["", "", "", "", "РАЗОМ:", t["total"]]]
            tbl = document.add_table(rows=len(rows), cols=len(t["headers"]))
            tbl.style = "Table Grid"
            for r_idx, row_data in enumerate(rows):
                row = tbl.rows[r_idx]
                for c_idx, cell_text in enumerate(row_data):
                    cell = row.cells[c_idx]
                    cell.text = str(cell_text)
                    run_cell = (cell.paragraphs[0].runs[0]
                                if cell.paragraphs[0].runs
                                else cell.paragraphs[0].add_run(str(cell_text)))
                    run_cell.font.size = Pt(9)
                    if r_idx == 0:
                        run_cell.bold = True
                    _set_run_font(run_cell)

    document.add_paragraph()
    sig_table = document.add_table(rows=2, cols=2)
    for i, label in enumerate(["ВИКОНАВЕЦЬ / ЗАЯВНИК", "ЗАМОВНИК / ОТРИМУВАЧ"]):
        cell = sig_table.rows[0].cells[i]
        cell.text = label
        run_h = cell.paragraphs[0].runs[0]
        run_h.bold = True
        run_h.font.size = Pt(10)
        _set_run_font(run_h)
    for i in range(2):
        cell = sig_table.rows[1].cells[i]
        cell.text = "________________  /_____________/"
        run_s = cell.paragraphs[0].runs[0]
        run_s.font.size = Pt(10)
        _set_run_font(run_s)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ----------------------------------------------------------------------------
# FLASK ROUTES
# ----------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/generate", methods=["POST"])
def generate():
    data     = request.get_json(force=True)
    doc_type = data.get("doc_type", "contract")
    fmt      = data.get("format", "pdf").lower()
    fields   = data.get("fields", {})

    if doc_type not in DOCUMENT_BUILDERS:
        return jsonify({"error": "Unknown document type: " + doc_type}), 400
    if fmt not in ("pdf", "docx"):
        return jsonify({"error": "Format must be 'pdf' or 'docx'"}), 400

    doc_data = DOCUMENT_BUILDERS[doc_type](fields)

    if fmt == "pdf":
        file_bytes = generate_pdf(doc_data)
        mimetype   = "application/pdf"
        ext        = "pdf"
    else:
        file_bytes = generate_docx(doc_data)
        mimetype   = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext        = "docx"

    filename = "%s_%s.%s" % (doc_type, fields.get("doc_number", uuid.uuid4().hex[:6]), ext)
    filepath = os.path.join(OUTPUT_DIR, filename)
    with open(filepath, "wb") as f:
        f.write(file_bytes)

    return send_file(
        io.BytesIO(file_bytes),
        mimetype=mimetype,
        as_attachment=True,
        download_name=filename,
    )


if __name__ == "__main__":
    app.run(debug=True, port=5000)
